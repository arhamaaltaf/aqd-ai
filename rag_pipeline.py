import json
import os
import pickle
import re
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, List, Optional


BASE_DIR = Path(__file__).resolve().parent

try:
    from dotenv import load_dotenv

    load_dotenv(BASE_DIR / ".env")
except Exception:
    pass

PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "aqdai1500")

embedding_model = None
reranker = None
llm = None
index = None
bm25_retriever = None
models_loaded = False


def environment_status() -> Dict[str, bool]:
    return {
        "groq_api_key": bool(os.getenv("GROQ_API_KEY")),
        "pinecone_api_key": bool(os.getenv("PINECONE_API_KEY")),
        "chunks_backup": (BASE_DIR / "chunks_backup.pkl").exists(),
    }


def load_models() -> None:
    global embedding_model, reranker, llm, index, bm25_retriever, models_loaded

    if models_loaded:
        return

    groq_api_key = os.getenv("GROQ_API_KEY")
    if not groq_api_key:
        raise RuntimeError("GROQ_API_KEY is not configured.")

    from langchain_community.retrievers import BM25Retriever
    from langchain_core.documents import Document
    from langchain_groq import ChatGroq
    from sentence_transformers import CrossEncoder, SentenceTransformer

    embedding_model = SentenceTransformer("BAAI/bge-base-en-v1.5")
    reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
    llm = ChatGroq(
        model=os.getenv("GROQ_MODEL", "llama-3.1-8b-instant"),
        api_key=groq_api_key,
        temperature=0.2,
        max_tokens=2500,
    )

    chunks_path = BASE_DIR / "chunks_backup.pkl"
    if not chunks_path.exists():
        raise RuntimeError("chunks_backup.pkl was not found in the project root.")

    with chunks_path.open("rb") as f:
        data = pickle.load(f)

    docs = [
        Document(page_content=chunk, metadata=metadata)
        for chunk, metadata in zip(data["chunks"], data["chunks_metadata"])
    ]
    bm25_retriever = BM25Retriever.from_documents(docs)
    bm25_retriever.k = 5

    pinecone_api_key = os.getenv("PINECONE_API_KEY")
    if pinecone_api_key:
        from pinecone import Pinecone

        pc = Pinecone(api_key=pinecone_api_key)
        index = pc.Index(PINECONE_INDEX_NAME)

    models_loaded = True


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    import fitz

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = fitz.open(tmp_path)
        text = ""
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            text += f"\n\n--- Page {page_num} ---\n\n{page_text}"
        doc.close()
        return text
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def split_user_contract(contract_text: str) -> List[str]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    numbered_headings = list(re.finditer(r"(?m)^\s*\d+\.\s+[^\n]+", contract_text))
    if numbered_headings:
        clause_chunks = []
        for index, match in enumerate(numbered_headings):
            start = match.start()
            end = numbered_headings[index + 1].start() if index + 1 < len(numbered_headings) else len(contract_text)
            chunk = contract_text[start:end].strip()
            chunk = re.sub(r"\n{3,}", "\n\n", chunk)
            if len(chunk) >= 50:
                clause_chunks.append(chunk)

        if clause_chunks:
            return clause_chunks

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=150,
        separators=[
            "\n\n",
            "\nWHEREAS",
            "\nNOW THEREFORE",
            "\nIN WITNESS",
            "\n",
            ". ",
            " ",
        ],
    )

    return [chunk.strip() for chunk in splitter.split_text(contract_text) if len(chunk.strip()) >= 80]


SHARIAH_RELEVANCE_TERMS = (
    "interest",
    "riba",
    "late payment",
    "penalty",
    "penalties",
    "compounding",
    "compounded",
    "finance charge",
    "outstanding amount",
    "default",
    "acceleration",
    "loan",
    "debt",
    "credit",
    "mark-up",
    "markup",
    "rent revision",
    "sole discretion",
    "without prior notice",
    "insurance",
    "takaful",
    "speculative",
    "market price",
    "uncertain",
    "discretion",
    "profit",
    "loss",
    "fees",
    "charges",
    "damages",
    "indemnity",
    "liability",
    "subject matter",
    "price",
    "payment",
    "delivery",
)


def is_obviously_shariah_relevant(chunk: str) -> bool:
    normalized = chunk.lower()
    return any(term in normalized for term in SHARIAH_RELEVANCE_TERMS)


def filter_relevant_chunks_with_llm(chunks: List[str], batch_size: int = 8) -> List[str]:
    relevant_chunks: List[str] = []
    llm_candidates: List[str] = []

    for chunk in chunks:
        if is_obviously_shariah_relevant(chunk):
            relevant_chunks.append(chunk)
        else:
            llm_candidates.append(chunk)

    if not llm_candidates:
        return relevant_chunks

    load_models()

    for start in range(0, len(llm_candidates), batch_size):
        batch = llm_candidates[start : start + batch_size]
        numbered_chunks = "\n\n".join(f"[{i + 1}] {chunk}" for i, chunk in enumerate(batch))

        prompt = f"""
You are a Shariah contract screening assistant.

Identify which contract chunks may require Shariah compliance analysis.

Mark a chunk as relevant if it contains anything related to:
- Payment terms, price, profit, loss, fees, installments, deposits, refunds
- Interest, late payment charges, penalties, compounding, finance charges
- Loans, debt, credit, mark-up, default
- Uncertainty in subject matter, price, delivery, obligations, or deliverables
- Speculation, gambling, outcome-based payments, market-linked returns
- Unfair or one-sided rights
- Risk transfer, indemnity, liability, damages
- Insurance or Takaful
- Haram goods or services
- Any other Islamic contract law or Shariah compliance concern

Do not select chunks that are only names, addresses, signatures, formatting,
generic boilerplate, notices, governing law, inspection, or assignment unless
they contain a real financial, risk, uncertainty, insurance, or compliance concern.

CONTRACT CHUNKS:
{numbered_chunks}

Reply with ONLY comma-separated numbers. If none are relevant, reply only NONE.
"""

        result = invoke_llm_with_retry(prompt).strip()
        if result.upper().startswith("NONE"):
            continue

        first_line = result.splitlines()[0].strip()
        selected_indexes = []
        for token in first_line.split(","):
            match = re.match(r"^\s*(\d+)\s*$", token)
            if match:
                selected_indexes.append(int(match.group(1)) - 1)

        if not selected_indexes:
            relevant_chunks.extend(batch)
            continue

        for idx in selected_indexes:
            if 0 <= idx < len(batch):
                relevant_chunks.append(batch[idx])

    return relevant_chunks


def semantic_search(query: str, k: int = 5) -> List[Any]:
    if index is None:
        return []

    from langchain_core.documents import Document

    query_vector = embedding_model.encode(query).tolist()
    results = index.query(vector=query_vector, top_k=k, include_metadata=True)
    return [
        Document(
            page_content=match["metadata"].get("text", ""),
            metadata={
                "source": match["metadata"].get("source", "Unknown"),
                "page": match["metadata"].get("page"),
                "score": match.get("score"),
            },
        )
        for match in results.get("matches", [])
        if match.get("metadata", {}).get("text")
    ]


def hybrid_search(query: str, k: int = 8) -> List[Any]:
    load_models()

    bm25_results = bm25_retriever.invoke(query) if bm25_retriever is not None else []
    semantic_results = semantic_search(query, k=k)

    scores: Dict[str, float] = {}
    all_docs: Dict[str, Any] = {}

    for rank, doc in enumerate(bm25_results):
        doc_id = doc.page_content[:150]
        scores[doc_id] = scores.get(doc_id, 0.0) + 1 / (rank + 60)
        all_docs[doc_id] = doc

    for rank, doc in enumerate(semantic_results):
        doc_id = doc.page_content[:150]
        scores[doc_id] = scores.get(doc_id, 0.0) + 1 / (rank + 60)
        all_docs[doc_id] = doc

    sorted_docs = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [all_docs[doc_id] for doc_id, _ in sorted_docs[:k]]


def get_shariah_context_for_contract_chunk(contract_chunk: str, k: int = 5) -> Dict[str, Any]:
    retrieved_docs = hybrid_search(contract_chunk, k=6)
    if not retrieved_docs:
        return {"context": "", "sources": []}

    pairs = [[contract_chunk, doc.page_content] for doc in retrieved_docs]
    scores = reranker.predict(pairs)
    ranked_docs = sorted(zip(scores, retrieved_docs), key=lambda x: x[0], reverse=True)
    top_docs = [doc for _, doc in ranked_docs[:k]]

    context = "\n\n---\n\n".join(
        f"Source: {doc.metadata.get('source', 'Unknown')}, Page: {doc.metadata.get('page', 'N/A')}\n{doc.page_content[:900]}"
        for doc in top_docs
    )
    sources = [
        {
            "source": doc.metadata.get("source", "Unknown"),
            "page": doc.metadata.get("page"),
            "content": doc.page_content[:350],
        }
        for doc in top_docs
    ]
    return {"context": context, "sources": sources}


def _extract_json(text: str) -> Optional[Dict[str, Any]]:
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group())
    except json.JSONDecodeError:
        return None


def invoke_llm_with_retry(prompt: str, attempts: int = 3) -> str:
    load_models()

    last_error: Optional[Exception] = None
    for attempt in range(attempts):
        try:
            return llm.invoke(prompt).content
        except Exception as exc:
            last_error = exc
            message = str(exc)
            if "rate_limit" not in message.lower() and "rate limit" not in message.lower() and "429" not in message:
                raise

            wait_seconds = 20
            match = re.search(r"try again in ([0-9.]+)s", message, re.IGNORECASE)
            if match:
                wait_seconds = min(35, max(5, int(float(match.group(1))) + 2))

            if attempt < attempts - 1:
                time.sleep(wait_seconds)

    if last_error:
        raise last_error
    raise RuntimeError("LLM call failed.")


def analyze_contract_chunk(contract_chunk: str, shariah_context: str) -> Dict[str, Any]:
    load_models()

    prompt = f"""
You are Aqd AI, a Shariah compliance assistant for contracts.

Analyze the user's contract chunk using only the Shariah sources below.

SHARIAH SOURCES:
{shariah_context}

USER CONTRACT CHUNK:
\"\"\"{contract_chunk}\"\"\"

Check for:
- Riba: interest, time-based increase on debt, compounding, late payment interest
- Gharar: excessive uncertainty, vague subject matter, unknown price, unclear delivery
- Maysir: gambling, speculation, zero-sum risk, outcome-based uncertainty
- Zulm: unfairness, oppression, one-sided rights, unfair risk allocation
- Invalid contract elements, non-compliant penalties, insurance issues, haram subject matter

Return ONLY valid JSON with this exact shape:
{{
  "verdict": "Compliant" | "Non-Compliant" | "Requires Scholarly Review",
  "issue_type": "Riba" | "Gharar" | "Maysir" | "Zulm" | "Invalid Contract Element" | "Penalty Issue" | "Insurance Issue" | "Other" | "None",
  "explanation": "plain English explanation",
  "principle_name": "short Islamic finance principle name",
  "principle_description": "one sentence principle description",
  "source_basis": "retrieved source name or short basis",
  "suggested_alternative": "Shariah-compliant improvement if needed, otherwise empty string"
}}

Do not invent AAOIFI standard numbers. If a late payment penalty is needed, say it should not become profit for the creditor and may be routed to charity, while documented actual losses may be claimed where allowed.
"""

    raw = invoke_llm_with_retry(prompt)
    parsed = _extract_json(raw)
    if parsed is None:
        return {
            "verdict": "Requires Scholarly Review",
            "issue_type": "Other",
            "explanation": raw.strip(),
            "principle_name": "Shariah Review Required",
            "principle_description": "A qualified Shariah reviewer should assess this clause.",
            "source_basis": "Retrieved Islamic finance context",
            "suggested_alternative": "",
        }
    return parsed


def generate_overall_contract_summary(report_items: List[Dict[str, Any]]) -> str:
    if not report_items:
        return "No Shariah-relevant clauses were found in the uploaded contract. This is an educational analysis, not a fatwa."

    non_compliant = [item for item in report_items if item.get("type") != "COMPLIANT"]
    issue_types = sorted({item.get("type", "OTHER") for item in non_compliant})

    if non_compliant:
        verdict = "Non-Compliant" if any(item.get("type") in {"RIBA", "MAYSIR"} for item in non_compliant) else "Requires Scholarly Review"
        issue_text = ", ".join(issue_types)
        serious = non_compliant[0].get("principleName", "the flagged clause")
        return (
            f"Overall verdict: {verdict}. Main issues flagged: {issue_text}. "
            f"The most serious concern appears to be {serious}. "
            "Review the highlighted clauses with a qualified Shariah scholar and revise the contract before signing. "
            "This is an educational analysis, not a fatwa."
        )

    return (
        "Overall verdict: Compliant based on the clauses reviewed. No major Shariah concern was flagged by the scanner. "
        "A final review by a qualified Shariah scholar is still recommended before signing. "
        "This is an educational analysis, not a fatwa."
    )


def normalize_issue_type(issue_type: str, verdict: str) -> str:
    normalized = (issue_type or "").strip().upper()
    verdict_normalized = (verdict or "").strip().upper()

    if verdict_normalized == "COMPLIANT" or normalized in {"NONE", "COMPLIANT"}:
        return "COMPLIANT"
    if "RIBA" in normalized or "PENALTY" in normalized:
        return "RIBA"
    if "GHARAR" in normalized or "INVALID" in normalized or "UNCERTAINTY" in normalized:
        return "GHARAR"
    if "MAYSIR" in normalized or "SPECULATION" in normalized:
        return "MAYSIR"
    return "GHARAR"


def analyze_contract_text(contract_text: str) -> Dict[str, Any]:
    load_models()

    if not contract_text.strip():
        raise ValueError("No contract text was provided.")

    all_chunks = split_user_contract(contract_text)
    if not all_chunks:
        raise ValueError("The contract text is too short or could not be split into clauses.")

    relevant_chunks = filter_relevant_chunks_with_llm(all_chunks)
    findings: List[Dict[str, Any]] = []
    all_sources: List[Dict[str, Any]] = []

    for idx, chunk in enumerate(relevant_chunks, start=1):
        context_payload = get_shariah_context_for_contract_chunk(chunk, k=3)
        analysis = analyze_contract_chunk(chunk, context_payload["context"])
        finding_type = normalize_issue_type(analysis.get("issue_type", ""), analysis.get("verdict", ""))

        sources = context_payload["sources"]
        if sources:
            all_sources.extend(sources)

        source_text = analysis.get("source_basis") or (
            f"{sources[0]['source']}, p. {sources[0].get('page')}" if sources else "Retrieved Islamic finance context"
        )
        suggested = analysis.get("suggested_alternative", "").strip()
        explanation = analysis.get("explanation", "").strip()
        if suggested:
            explanation = f"{explanation}\n\nSuggested alternative: {suggested}"

        findings.append(
            {
                "id": idx,
                "type": finding_type,
                "verdict": analysis.get("verdict", ""),
                "clauseText": chunk,
                "explanation": explanation,
                "principleName": analysis.get("principle_name", "Islamic Contract Principle"),
                "principleDescription": analysis.get("principle_description", ""),
                "source": source_text,
                "suggestedAlternative": suggested,
            }
        )

    violations = sum(1 for f in findings if f["type"] in {"RIBA", "MAYSIR"})
    uncertain = sum(1 for f in findings if f["type"] == "GHARAR")
    compliant = sum(1 for f in findings if f["type"] == "COMPLIANT")
    summary_text = generate_overall_contract_summary(findings)

    unique_sources = []
    seen = set()
    for source in all_sources:
        key = (source.get("source"), source.get("page"), source.get("content"))
        if key not in seen:
            seen.add(key)
            unique_sources.append(source)

    return {
        "contractText": contract_text,
        "findings": findings,
        "summary": {
            "violations": violations,
            "uncertain": uncertain,
            "compliant": compliant,
            "text": summary_text,
            "chunksScanned": len(all_chunks),
            "relevantChunks": len(relevant_chunks),
        },
        "sources": unique_sources[:20],
    }
